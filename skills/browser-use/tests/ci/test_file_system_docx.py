"""Tests for DOCX file support in the FileSystem."""

from pathlib import Path

import pytest

from browser_use.filesystem.file_system import (
	DocxFile,
	FileSystem,
)


class TestDocxFile:
	"""Test DOCX file operations."""

	@pytest.mark.asyncio
	async def test_create_docx_file(self, tmp_path: Path):
		"""Test creating a DOCX file."""
		fs = FileSystem(tmp_path)
		content = """# Heading 1
## Heading 2
### Heading 3
Regular paragraph text.

Another paragraph."""

		result = await fs.write_file('test.docx', content)
		assert 'successfully' in result.lower()
		assert 'test.docx' in fs.list_files()

	@pytest.mark.asyncio
	async def test_read_docx_file_internal(self, tmp_path: Path):
		"""Test reading internal DOCX file."""
		fs = FileSystem(tmp_path)
		content = """# Title
Some content here."""

		await fs.write_file('test.docx', content)
		result = await fs.read_file('test.docx')

		assert 'test.docx' in result
		assert 'Title' in result or 'content' in result

	@pytest.mark.asyncio
	async def test_read_docx_file_external(self, tmp_path: Path):
		"""Test reading external DOCX file."""
		from docx import Document

		# Create an external DOCX file
		external_file = tmp_path / 'external.docx'
		doc = Document()
		doc.add_heading('Test Heading', level=1)
		doc.add_paragraph('Test paragraph content.')
		doc.save(str(external_file))

		fs = FileSystem(tmp_path / 'workspace')
		structured_result = await fs.read_file_structured(str(external_file), external_file=True)

		assert 'message' in structured_result
		assert 'Test Heading' in structured_result['message']
		assert 'Test paragraph content' in structured_result['message']

	def test_docx_file_extension(self):
		"""Test DOCX file extension property."""
		docx_file = DocxFile(name='test')
		assert docx_file.extension == 'docx'
		assert docx_file.full_name == 'test.docx'

	@pytest.mark.asyncio
	async def test_docx_with_unicode_characters(self, tmp_path: Path):
		"""Test DOCX with unicode and emoji content."""
		fs = FileSystem(tmp_path)
		content = """# Unicode Test 🚀
Chinese: 你好世界
Arabic: مرحبا بالعالم
Emoji: 😀 👍 🎉"""

		result = await fs.write_file('unicode.docx', content)
		assert 'successfully' in result.lower()

		read_result = await fs.read_file('unicode.docx')
		assert 'Unicode Test' in read_result
		# Note: Emoji may not be preserved in all systems

	@pytest.mark.asyncio
	async def test_empty_docx_file(self, tmp_path: Path):
		"""Test creating an empty DOCX file."""
		fs = FileSystem(tmp_path)
		result = await fs.write_file('empty.docx', '')
		assert 'successfully' in result.lower()

	@pytest.mark.asyncio
	async def test_large_docx_file(self, tmp_path: Path):
		"""Test creating a large DOCX file."""
		fs = FileSystem(tmp_path)
		# Create content with 1000 lines
		lines = [f'Line {i}: This is a test line with some content.' for i in range(1000)]
		content = '\n'.join(lines)

		result = await fs.write_file('large.docx', content)
		assert 'successfully' in result.lower()

		# Verify it can be read back
		read_result = await fs.read_file('large.docx')
		assert 'Line 0:' in read_result
		assert 'Line 999:' in read_result

	@pytest.mark.asyncio
	async def test_corrupted_docx_file(self, tmp_path: Path):
		"""Test reading a corrupted DOCX file."""
		# Create a corrupted DOCX file
		external_file = tmp_path / 'corrupted.docx'
		external_file.write_bytes(b'This is not a valid DOCX file')

		fs = FileSystem(tmp_path / 'workspace')
		structured_result = await fs.read_file_structured(str(external_file), external_file=True)

		assert 'message' in structured_result
		assert 'error' in structured_result['message'].lower() or 'could not' in structured_result['message'].lower()

	@pytest.mark.asyncio
	async def test_docx_with_multiple_paragraphs(self, tmp_path: Path):
		"""Test DOCX with various paragraph styles."""
		fs = FileSystem(tmp_path)
		content = """# Main Title
## Subtitle
This is a regular paragraph.

This is another paragraph with some text.

### Section 3
Final paragraph here."""

		await fs.write_file('multi.docx', content)
		result = await fs.read_file('multi.docx')

		# Should contain all the text (headings converted to paragraphs)
		assert 'Main Title' in result
		assert 'Subtitle' in result
		assert 'regular paragraph' in result
		assert 'Final paragraph' in result


class TestFileSystemDocxIntegration:
	"""Integration tests for DOCX file type."""

	@pytest.mark.asyncio
	async def test_multiple_file_types_with_docx(self, tmp_path: Path):
		"""Test working with DOCX alongside other file types."""
		fs = FileSystem(tmp_path)

		# Create different file types
		await fs.write_file('doc.docx', '# Document\nContent here')
		await fs.write_file('data.json', '{"key": "value"}')
		await fs.write_file('notes.txt', 'Some notes')

		# Verify all files exist
		files = fs.list_files()
		assert 'doc.docx' in files
		assert 'data.json' in files
		assert 'notes.txt' in files
		assert 'todo.md' in files  # Default file

	@pytest.mark.asyncio
	async def test_file_system_state_with_docx(self, tmp_path: Path):
		"""Test FileSystem state serialization with DOCX files."""
		fs = FileSystem(tmp_path)

		# Create files
		await fs.write_file('test.docx', '# Title\nContent')
		await fs.write_file('data.txt', 'Some text')

		# Get state
		state = fs.get_state()
		assert 'test.docx' in state.files
		assert 'data.txt' in state.files

		# Restore from state
		fs2 = FileSystem.from_state(state)
		assert 'test.docx' in fs2.list_files()
		assert 'data.txt' in fs2.list_files()

	def test_allowed_extensions_include_docx(self, tmp_path: Path):
		"""Test that DOCX is in allowed extensions."""
		fs = FileSystem(tmp_path)
		allowed = fs.get_allowed_extensions()

		assert 'docx' in allowed


if __name__ == '__main__':
	pytest.main([__file__, '-v'])
