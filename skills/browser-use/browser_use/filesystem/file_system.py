import asyncio
import base64
import os
import re
import shutil
from abc import ABC, abstractmethod
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

UNSUPPORTED_BINARY_EXTENSIONS = {
	'png',
	'jpg',
	'jpeg',
	'gif',
	'bmp',
	'svg',
	'webp',
	'ico',
	'mp3',
	'mp4',
	'wav',
	'avi',
	'mov',
	'zip',
	'tar',
	'gz',
	'rar',
	'exe',
	'bin',
	'dll',
	'so',
}


def _build_filename_error_message(file_name: str, supported_extensions: list[str]) -> str:
	"""Build a specific error message explaining why the filename was rejected and how to fix it."""
	base = os.path.basename(file_name)

	# Check for binary/image extension
	if '.' in base:
		_, ext = base.rsplit('.', 1)
		ext_lower = ext.lower()
		if ext_lower in UNSUPPORTED_BINARY_EXTENSIONS:
			return (
				f"Error: Cannot write binary/image file '{base}'. "
				f'The write_file tool only supports text-based files. '
				f'Supported extensions: {", ".join("." + e for e in supported_extensions)}. '
				f'For screenshots, the browser automatically captures them - do not try to save screenshots as files.'
			)
		if ext_lower not in supported_extensions:
			return (
				f"Error: Unsupported file extension '.{ext_lower}' in '{base}'. "
				f'Supported extensions: {", ".join("." + e for e in supported_extensions)}. '
				f'Please rename the file to use a supported extension.'
			)

	# No extension or no dot
	if '.' not in base:
		return (
			f"Error: Filename '{base}' has no extension. "
			f'Please add a supported extension: {", ".join("." + e for e in supported_extensions)}.'
		)

	return (
		f"Error: Invalid filename '{base}'. "
		f'Filenames must contain only letters, numbers, underscores, hyphens, dots, parentheses, and spaces. '
		f'Supported extensions: {", ".join("." + e for e in supported_extensions)}.'
	)


DEFAULT_FILE_SYSTEM_PATH = 'browseruse_agent_data'


class FileSystemError(Exception):
	"""Custom exception for file system operations that should be shown to LLM"""

	pass


class BaseFile(BaseModel, ABC):
	"""Base class for all file types"""

	name: str
	content: str = ''

	# --- Subclass must define this ---
	@property
	@abstractmethod
	def extension(self) -> str:
		"""File extension (e.g. 'txt', 'md')"""
		pass

	def write_file_content(self, content: str) -> None:
		"""Update internal content (formatted)"""
		self.update_content(content)

	def append_file_content(self, content: str) -> None:
		"""Append content to internal content"""
		self.update_content(self.content + content)

	# --- These are shared and implemented here ---

	def update_content(self, content: str) -> None:
		self.content = content

	def sync_to_disk_sync(self, path: Path) -> None:
		file_path = path / self.full_name
		file_path.write_text(self.content)

	async def sync_to_disk(self, path: Path) -> None:
		file_path = path / self.full_name
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: file_path.write_text(self.content))

	async def write(self, content: str, path: Path) -> None:
		self.write_file_content(content)
		await self.sync_to_disk(path)

	async def append(self, content: str, path: Path) -> None:
		self.append_file_content(content)
		await self.sync_to_disk(path)

	def read(self) -> str:
		return self.content

	@property
	def full_name(self) -> str:
		return f'{self.name}.{self.extension}'

	@property
	def get_size(self) -> int:
		return len(self.content)

	@property
	def get_line_count(self) -> int:
		return len(self.content.splitlines())


class MarkdownFile(BaseFile):
	"""Markdown file implementation"""

	@property
	def extension(self) -> str:
		return 'md'


class TxtFile(BaseFile):
	"""Plain text file implementation"""

	@property
	def extension(self) -> str:
		return 'txt'


class JsonFile(BaseFile):
	"""JSON file implementation"""

	@property
	def extension(self) -> str:
		return 'json'


class CsvFile(BaseFile):
	"""CSV file implementation"""

	@property
	def extension(self) -> str:
		return 'csv'


class JsonlFile(BaseFile):
	"""JSONL (JSON Lines) file implementation"""

	@property
	def extension(self) -> str:
		return 'jsonl'


class PdfFile(BaseFile):
	"""PDF file implementation"""

	@property
	def extension(self) -> str:
		return 'pdf'

	def sync_to_disk_sync(self, path: Path) -> None:
		# Lazy import reportlab
		from reportlab.lib.pagesizes import letter
		from reportlab.lib.styles import getSampleStyleSheet
		from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

		file_path = path / self.full_name
		try:
			# Create PDF document
			doc = SimpleDocTemplate(str(file_path), pagesize=letter)
			styles = getSampleStyleSheet()
			story = []

			# Convert markdown content to simple text and add to PDF
			# For basic implementation, we'll treat content as plain text
			# This avoids the AGPL license issue while maintaining functionality
			content_lines = self.content.split('\n')

			for line in content_lines:
				if line.strip():
					# Handle basic markdown headers
					if line.startswith('# '):
						para = Paragraph(line[2:], styles['Title'])
					elif line.startswith('## '):
						para = Paragraph(line[3:], styles['Heading1'])
					elif line.startswith('### '):
						para = Paragraph(line[4:], styles['Heading2'])
					else:
						para = Paragraph(line, styles['Normal'])
					story.append(para)
				else:
					story.append(Spacer(1, 6))

			doc.build(story)
		except Exception as e:
			raise FileSystemError(f"Error: Could not write to file '{self.full_name}'. {str(e)}")

	async def sync_to_disk(self, path: Path) -> None:
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: self.sync_to_disk_sync(path))


class DocxFile(BaseFile):
	"""DOCX file implementation"""

	@property
	def extension(self) -> str:
		return 'docx'

	def sync_to_disk_sync(self, path: Path) -> None:
		file_path = path / self.full_name
		try:
			from docx import Document

			doc = Document()

			# Convert content to DOCX paragraphs
			content_lines = self.content.split('\n')

			for line in content_lines:
				if line.strip():
					# Handle basic markdown headers
					if line.startswith('# '):
						doc.add_heading(line[2:], level=1)
					elif line.startswith('## '):
						doc.add_heading(line[3:], level=2)
					elif line.startswith('### '):
						doc.add_heading(line[4:], level=3)
					else:
						doc.add_paragraph(line)
				else:
					doc.add_paragraph()  # Empty paragraph for spacing

			doc.save(str(file_path))
		except Exception as e:
			raise FileSystemError(f"Error: Could not write to file '{self.full_name}'. {str(e)}")

	async def sync_to_disk(self, path: Path) -> None:
		with ThreadPoolExecutor() as executor:
			await asyncio.get_event_loop().run_in_executor(executor, lambda: self.sync_to_disk_sync(path))


class HtmlFile(BaseFile):
	"""HTML file implementation"""

	@property
	def extension(self) -> str:
		return 'html'


class XmlFile(BaseFile):
	"""XML file implementation"""

	@property
	def extension(self) -> str:
		return 'xml'


class FileSystemState(BaseModel):
	"""Serializable state of the file system"""

	files: dict[str, dict[str, Any]] = Field(default_factory=dict)  # full filename -> file data
	base_dir: str
	extracted_content_count: int = 0


class FileSystem:
	"""Enhanced file system with in-memory storage and multiple file type support"""

	def __init__(self, base_dir: str | Path, create_default_files: bool = True):
		# Handle the Path conversion before calling super().__init__
		self.base_dir = Path(base_dir) if isinstance(base_dir, str) else base_dir
		self.base_dir.mkdir(parents=True, exist_ok=True)

		# Create and use a dedicated subfolder for all operations
		self.data_dir = self.base_dir / DEFAULT_FILE_SYSTEM_PATH
		if self.data_dir.exists():
			# clean the data directory
			shutil.rmtree(self.data_dir)
		self.data_dir.mkdir(exist_ok=True)

		self._file_types: dict[str, type[BaseFile]] = {
			'md': MarkdownFile,
			'txt': TxtFile,
			'json': JsonFile,
			'jsonl': JsonlFile,
			'csv': CsvFile,
			'pdf': PdfFile,
			'docx': DocxFile,
			'html': HtmlFile,
			'xml': XmlFile,
		}

		self.files = {}
		if create_default_files:
			self.default_files = ['todo.md']
			self._create_default_files()

		self.extracted_content_count = 0

	def get_allowed_extensions(self) -> list[str]:
		"""Get allowed extensions"""
		return list(self._file_types.keys())

	def _get_file_type_class(self, extension: str) -> type[BaseFile] | None:
		"""Get the appropriate file class for an extension."""
		return self._file_types.get(extension.lower(), None)

	def _create_default_files(self) -> None:
		"""Create default results and todo files"""
		for full_filename in self.default_files:
			name_without_ext, extension = self._parse_filename(full_filename)
			file_class = self._get_file_type_class(extension)
			if not file_class:
				raise ValueError(f"Error: Invalid file extension '{extension}' for file '{full_filename}'.")

			file_obj = file_class(name=name_without_ext)
			self.files[full_filename] = file_obj  # Use full filename as key
			file_obj.sync_to_disk_sync(self.data_dir)

	def _is_valid_filename(self, file_name: str) -> bool:
		"""Check if filename matches the required pattern: name.extension

		Allows letters, numbers, underscores, hyphens, dots, parentheses, spaces, and Chinese characters
		in the name part, followed by a dot and a supported extension.
		"""
		extensions = '|'.join(self._file_types.keys())
		# Allow dots, spaces, parens in the name part - match everything up to the last dot
		pattern = rf'^[a-zA-Z0-9_\-\.\(\) \u4e00-\u9fff]+\.({extensions})$'
		file_name_base = os.path.basename(file_name)
		if not re.match(pattern, file_name_base):
			return False
		# Ensure the name part (before last dot) is non-empty
		name_part = file_name_base.rsplit('.', 1)[0]
		return len(name_part.strip()) > 0

	@staticmethod
	def sanitize_filename(file_name: str) -> str:
		"""Sanitize a filename by replacing/removing invalid characters.

		- Replaces spaces with hyphens
		- Removes characters that are not alphanumeric, underscore, hyphen, dot, parentheses, or Chinese
		- Preserves the extension
		- Collapses multiple consecutive hyphens
		"""
		base = os.path.basename(file_name)
		if '.' not in base:
			return base

		name_part, ext = base.rsplit('.', 1)
		# Replace spaces with hyphens
		name_part = name_part.replace(' ', '-')
		# Remove invalid characters (keep alphanumeric, underscore, hyphen, dot, parens, Chinese)
		name_part = re.sub(r'[^a-zA-Z0-9_\-\.\(\)\u4e00-\u9fff]', '', name_part)
		# Collapse multiple hyphens
		name_part = re.sub(r'-{2,}', '-', name_part)
		# Strip leading/trailing hyphens and dots
		name_part = name_part.strip('-.')

		if not name_part:
			name_part = 'file'

		return f'{name_part}.{ext.lower()}'

	def _resolve_filename(self, file_name: str) -> tuple[str, bool]:
		"""Resolve a filename, attempting sanitization if the original is invalid.

		Normalizes to basename first to prevent directory traversal (e.g. ../secret.md).

		Returns:
			(resolved_name, was_changed): The resolved filename and whether it differs from the input.
			If resolution fails, returns (basename, was_changed).
		"""
		base_name = os.path.basename(file_name)
		was_changed = base_name != file_name

		if self._is_valid_filename(base_name):
			return base_name, was_changed

		sanitized = self.sanitize_filename(base_name)
		if sanitized != base_name and self._is_valid_filename(sanitized):
			return sanitized, True

		return base_name, was_changed

	def _parse_filename(self, filename: str) -> tuple[str, str]:
		"""Parse filename into name and extension. Always check _is_valid_filename first."""
		name, extension = filename.rsplit('.', 1)
		return name, extension.lower()

	def get_dir(self) -> Path:
		"""Get the file system directory"""
		return self.data_dir

	def get_file(self, full_filename: str) -> BaseFile | None:
		"""Get a file object by full filename, trying sanitization if the name is invalid."""
		resolved, _ = self._resolve_filename(full_filename)
		if not self._is_valid_filename(resolved):
			return None

		# Use resolved filename as key
		return self.files.get(resolved)

	def list_files(self) -> list[str]:
		"""List all files in the system"""
		return [file_obj.full_name for file_obj in self.files.values()]

	def display_file(self, full_filename: str) -> str | None:
		"""Display file content using file-specific display method"""
		resolved, _ = self._resolve_filename(full_filename)
		if not self._is_valid_filename(resolved):
			return None

		file_obj = self.files.get(resolved)
		if not file_obj:
			return None

		return file_obj.read()

	async def read_file_structured(self, full_filename: str, external_file: bool = False) -> dict[str, Any]:
		"""Read file and return structured data including images if applicable.

		Returns:
			dict with keys:
				- 'message': str - The message to display
				- 'images': list[dict] | None - Image data if file is an image: [{"name": str, "data": base64_str}]
		"""
		result: dict[str, Any] = {'message': '', 'images': None}

		if external_file:
			try:
				try:
					_, extension = self._parse_filename(full_filename)
				except Exception:
					result['message'] = (
						f'Error: Invalid filename format {full_filename}. Must be alphanumeric with a supported extension.'
					)
					return result

				# Text-based extensions: derive from _file_types, excluding those with special readers
				_special_extensions = {'docx', 'pdf', 'jpg', 'jpeg', 'png'}
				text_extensions = [ext for ext in self._file_types if ext not in _special_extensions]

				if extension in text_extensions:
					import anyio

					async with await anyio.open_file(full_filename, 'r') as f:
						content = await f.read()
						result['message'] = f'Read from file {full_filename}.\n<content>\n{content}\n</content>'
						return result

				elif extension == 'docx':
					from docx import Document

					doc = Document(full_filename)
					content = '\n'.join([para.text for para in doc.paragraphs])
					result['message'] = f'Read from file {full_filename}.\n<content>\n{content}\n</content>'
					return result

				elif extension == 'pdf':
					import pypdf

					reader = pypdf.PdfReader(full_filename)
					num_pages = len(reader.pages)
					MAX_CHARS = 60000  # character-based limit

					# Extract text from all pages with page markers
					page_texts: list[tuple[int, str]] = []
					total_chars = 0
					for i, page in enumerate(reader.pages, 1):
						text = page.extract_text() or ''
						page_texts.append((i, text))
						total_chars += len(text)

					# If small enough, return everything
					if total_chars <= MAX_CHARS:
						content_parts = []
						for page_num, text in page_texts:
							if text.strip():
								content_parts.append(f'--- Page {page_num} ---\n{text}')
						extracted_text = '\n\n'.join(content_parts)
						result['message'] = (
							f'Read from file {full_filename} ({num_pages} pages, {total_chars:,} chars).\n'
							f'<content>\n{extracted_text}\n</content>'
						)
						return result

					# Large PDF - use search to prioritize pages with distinctive content
					import math
					import re

					# Extract words from each page and count which pages they appear on
					word_to_pages: dict[str, set[int]] = {}
					page_words: dict[int, set[str]] = {}

					for page_num, text in page_texts:
						# Extract words (lowercase, 4+ chars to filter noise)
						words = set(re.findall(r'\b[a-zA-Z]{4,}\b', text.lower()))
						page_words[page_num] = words
						for word in words:
							if word not in word_to_pages:
								word_to_pages[word] = set()
							word_to_pages[word].add(page_num)

					# Score pages using inverse document frequency (IDF)
					# words appearing on fewer pages get higher weight
					page_scores: dict[int, float] = {}
					for page_num, words in page_words.items():
						score = 0.0
						for word in words:
							pages_with_word = len(word_to_pages[word])
							# IDF: log(total_pages / pages_with_word) - higher for rarer words
							score += math.log(num_pages / pages_with_word)
						page_scores[page_num] = score

					# Sort pages by score (highest first), always include page 1
					sorted_pages = sorted(page_scores.items(), key=lambda x: -x[1])
					priority_pages = [1]
					for page_num, _ in sorted_pages:
						if page_num not in priority_pages:
							priority_pages.append(page_num)

					# Add remaining pages in order (for pages with no distinctive content)
					for page_num, _ in page_texts:
						if page_num not in priority_pages:
							priority_pages.append(page_num)

					# Build content from prioritized pages, respecting char limit
					content_parts = []
					chars_used = 0
					pages_included = []

					# First pass: add pages in priority order
					for page_num in priority_pages:
						text = page_texts[page_num - 1][1]
						if not text.strip():
							continue
						page_header = f'--- Page {page_num} ---\n'
						truncation_suffix = '\n[...truncated]'
						remaining = MAX_CHARS - chars_used
						# Need room for header + suffix + at least some content
						min_useful = len(page_header) + len(truncation_suffix) + 50
						if remaining < min_useful:
							break  # no room left for meaningful content
						page_content = page_header + text
						if len(page_content) > remaining:
							# Truncate page to fit remaining budget exactly
							page_content = page_content[: remaining - len(truncation_suffix)] + truncation_suffix
						content_parts.append((page_num, page_content))
						chars_used += len(page_content)
						pages_included.append(page_num)
						if chars_used >= MAX_CHARS:
							break

					# Sort included pages by page number for readability
					content_parts.sort(key=lambda x: x[0])
					extracted_text = '\n\n'.join(part for _, part in content_parts)

					pages_not_shown = num_pages - len(pages_included)
					if pages_not_shown > 0:
						skipped = [p for p in range(1, num_pages + 1) if p not in pages_included]
						truncation_note = (
							f'\n\n[Showing {len(pages_included)} of {num_pages} pages. '
							f'Skipped pages: {skipped[:10]}{"..." if len(skipped) > 10 else ""}. '
							f'Use read_long_content with a specific goal to find relevant sections.]'
						)
					else:
						truncation_note = ''

					result['message'] = (
						f'Read from file {full_filename} ({num_pages} pages, {total_chars:,} chars total).\n'
						f'<content>\n{extracted_text}{truncation_note}\n</content>'
					)
					return result

				elif extension in ['jpg', 'jpeg', 'png']:
					import anyio

					# Read image file and convert to base64
					async with await anyio.open_file(full_filename, 'rb') as f:
						img_data = await f.read()

					base64_str = base64.b64encode(img_data).decode('utf-8')

					result['message'] = f'Read image file {full_filename}.'
					result['images'] = [{'name': os.path.basename(full_filename), 'data': base64_str}]
					return result

				else:
					result['message'] = f'Error: Cannot read file {full_filename} as {extension} extension is not supported.'
					return result

			except FileNotFoundError:
				result['message'] = f"Error: File '{full_filename}' not found."
				return result
			except PermissionError:
				result['message'] = f"Error: Permission denied to read file '{full_filename}'."
				return result
			except Exception as e:
				result['message'] = f"Error: Could not read file '{full_filename}'. {str(e)}"
				return result

		# For internal files, only non-image types are supported
		resolved, was_sanitized = self._resolve_filename(full_filename)
		if not self._is_valid_filename(resolved):
			result['message'] = _build_filename_error_message(full_filename, self.get_allowed_extensions())
			return result

		file_obj = self.files.get(resolved)
		if not file_obj:
			if was_sanitized:
				result['message'] = f"File '{resolved}' not found. (Filename was auto-corrected from '{full_filename}')"
			else:
				result['message'] = f"File '{full_filename}' not found."
			return result

		try:
			content = file_obj.read()
			sanitize_note = f"Note: filename was auto-corrected from '{full_filename}' to '{resolved}'. " if was_sanitized else ''
			result['message'] = f'{sanitize_note}Read from file {resolved}.\n<content>\n{content}\n</content>'
			return result
		except FileSystemError as e:
			result['message'] = str(e)
			return result
		except Exception as e:
			result['message'] = f"Error: Could not read file '{full_filename}'. {str(e)}"
			return result

	async def read_file(self, full_filename: str, external_file: bool = False) -> str:
		"""Read file content using file-specific read method and return appropriate message to LLM.

		Note: For image files, use read_file_structured() to get image data.
		"""
		result = await self.read_file_structured(full_filename, external_file)
		return result['message']

	async def write_file(self, full_filename: str, content: str) -> str:
		"""Write content to file using file-specific write method"""
		original_filename = full_filename
		resolved, was_sanitized = self._resolve_filename(full_filename)
		if not self._is_valid_filename(resolved):
			return _build_filename_error_message(full_filename, self.get_allowed_extensions())
		full_filename = resolved

		try:
			name_without_ext, extension = self._parse_filename(full_filename)
			file_class = self._get_file_type_class(extension)
			if not file_class:
				raise ValueError(f"Error: Invalid file extension '{extension}' for file '{full_filename}'.")

			# Create or get existing file using full filename as key
			if full_filename in self.files:
				file_obj = self.files[full_filename]
			else:
				file_obj = file_class(name=name_without_ext)
				self.files[full_filename] = file_obj  # Use full filename as key

			# Use file-specific write method
			await file_obj.write(content, self.data_dir)
			sanitize_note = f" (auto-corrected from '{original_filename}')" if was_sanitized else ''
			return f'Data written to file {full_filename} successfully.{sanitize_note}'
		except FileSystemError as e:
			return str(e)
		except Exception as e:
			return f"Error: Could not write to file '{full_filename}'. {str(e)}"

	async def append_file(self, full_filename: str, content: str) -> str:
		"""Append content to file using file-specific append method"""
		original_filename = full_filename
		resolved, was_sanitized = self._resolve_filename(full_filename)
		if not self._is_valid_filename(resolved):
			return _build_filename_error_message(full_filename, self.get_allowed_extensions())
		full_filename = resolved

		file_obj = self.files.get(full_filename)
		if not file_obj:
			if was_sanitized:
				return f"File '{full_filename}' not found. (Filename was auto-corrected from '{original_filename}')"
			return f"File '{full_filename}' not found."

		try:
			await file_obj.append(content, self.data_dir)
			sanitize_note = f" (auto-corrected from '{original_filename}')" if was_sanitized else ''
			return f'Data appended to file {full_filename} successfully.{sanitize_note}'
		except FileSystemError as e:
			return str(e)
		except Exception as e:
			return f"Error: Could not append to file '{full_filename}'. {str(e)}"

	async def replace_file_str(self, full_filename: str, old_str: str, new_str: str) -> str:
		"""Replace old_str with new_str in file_name"""
		original_filename = full_filename
		resolved, was_sanitized = self._resolve_filename(full_filename)
		if not self._is_valid_filename(resolved):
			return _build_filename_error_message(full_filename, self.get_allowed_extensions())
		full_filename = resolved

		if not old_str:
			return 'Error: Cannot replace empty string. Please provide a non-empty string to replace.'

		file_obj = self.files.get(full_filename)
		if not file_obj:
			if was_sanitized:
				return f"File '{full_filename}' not found. (Filename was auto-corrected from '{original_filename}')"
			return f"File '{full_filename}' not found."

		try:
			content = file_obj.read()
			content = content.replace(old_str, new_str)
			await file_obj.write(content, self.data_dir)
			sanitize_note = f" (auto-corrected from '{original_filename}')" if was_sanitized else ''
			return f'Successfully replaced all occurrences of "{old_str}" with "{new_str}" in file {full_filename}{sanitize_note}'
		except FileSystemError as e:
			return str(e)
		except Exception as e:
			return f"Error: Could not replace string in file '{full_filename}'. {str(e)}"

	async def save_extracted_content(self, content: str) -> str:
		"""Save extracted content to a numbered file"""
		initial_filename = f'extracted_content_{self.extracted_content_count}'
		extracted_filename = f'{initial_filename}.md'
		file_obj = MarkdownFile(name=initial_filename)
		await file_obj.write(content, self.data_dir)
		self.files[extracted_filename] = file_obj
		self.extracted_content_count += 1
		return extracted_filename

	def describe(self) -> str:
		"""List all files with their content information using file-specific display methods"""
		DISPLAY_CHARS = 400
		description = ''

		for file_obj in self.files.values():
			# Skip todo.md from description
			if file_obj.full_name == 'todo.md':
				continue

			content = file_obj.read()

			# Handle empty files
			if not content:
				description += f'<file>\n{file_obj.full_name} - [empty file]\n</file>\n'
				continue

			lines = content.splitlines()
			line_count = len(lines)

			# For small files, display the entire content
			whole_file_description = (
				f'<file>\n{file_obj.full_name} - {line_count} lines\n<content>\n{content}\n</content>\n</file>\n'
			)
			if len(content) < int(1.5 * DISPLAY_CHARS):
				description += whole_file_description
				continue

			# For larger files, display start and end previews
			half_display_chars = DISPLAY_CHARS // 2

			# Get start preview
			start_preview = ''
			start_line_count = 0
			chars_count = 0
			for line in lines:
				if chars_count + len(line) + 1 > half_display_chars:
					break
				start_preview += line + '\n'
				chars_count += len(line) + 1
				start_line_count += 1

			# Get end preview
			end_preview = ''
			end_line_count = 0
			chars_count = 0
			for line in reversed(lines):
				if chars_count + len(line) + 1 > half_display_chars:
					break
				end_preview = line + '\n' + end_preview
				chars_count += len(line) + 1
				end_line_count += 1

			# Calculate lines in between
			middle_line_count = line_count - start_line_count - end_line_count
			if middle_line_count <= 0:
				description += whole_file_description
				continue

			start_preview = start_preview.strip('\n').rstrip()
			end_preview = end_preview.strip('\n').rstrip()

			# Format output
			if not (start_preview or end_preview):
				description += f'<file>\n{file_obj.full_name} - {line_count} lines\n<content>\n{middle_line_count} lines...\n</content>\n</file>\n'
			else:
				description += f'<file>\n{file_obj.full_name} - {line_count} lines\n<content>\n{start_preview}\n'
				description += f'... {middle_line_count} more lines ...\n'
				description += f'{end_preview}\n'
				description += '</content>\n</file>\n'

		return description.strip('\n')

	def get_todo_contents(self) -> str:
		"""Get todo file contents"""
		todo_file = self.get_file('todo.md')
		return todo_file.read() if todo_file else ''

	def get_state(self) -> FileSystemState:
		"""Get serializable state of the file system"""
		files_data = {}
		for full_filename, file_obj in self.files.items():
			files_data[full_filename] = {'type': file_obj.__class__.__name__, 'data': file_obj.model_dump()}

		return FileSystemState(
			files=files_data, base_dir=str(self.base_dir), extracted_content_count=self.extracted_content_count
		)

	def nuke(self) -> None:
		"""Delete the file system directory"""
		shutil.rmtree(self.data_dir)

	@classmethod
	def from_state(cls, state: FileSystemState) -> 'FileSystem':
		"""Restore file system from serializable state at the exact same location"""
		# Create file system without default files
		fs = cls(base_dir=Path(state.base_dir), create_default_files=False)
		fs.extracted_content_count = state.extracted_content_count

		# Restore all files
		for full_filename, file_data in state.files.items():
			file_type = file_data['type']
			file_info = file_data['data']

			# Create the appropriate file object based on type
			file_type_map: dict[str, type[BaseFile]] = {
				'MarkdownFile': MarkdownFile,
				'TxtFile': TxtFile,
				'JsonFile': JsonFile,
				'JsonlFile': JsonlFile,
				'CsvFile': CsvFile,
				'PdfFile': PdfFile,
				'DocxFile': DocxFile,
				'HtmlFile': HtmlFile,
				'XmlFile': XmlFile,
			}

			file_class = file_type_map.get(file_type)
			if not file_class:
				# Skip unknown file types
				continue
			file_obj = file_class(**file_info)

			# Add to files dict and sync to disk
			fs.files[full_filename] = file_obj
			file_obj.sync_to_disk_sync(fs.data_dir)

		return fs
